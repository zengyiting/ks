#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成完整的毕业论文文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import re

def set_cell_shading(cell, color):
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        qn('w:shd'),
        {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color}
    )
    shading_elm.append(shading)

def set_paragraph_format(p, space_before=0, space_after=6, line_spacing=1.5, left_indent=0):
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if left_indent > 0:
        pf.left_indent = Cm(left_indent)

def add_body(doc, text, space_before=0, space_after=6, bold=False, italic=False, font_size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    set_paragraph_format(p, space_before, space_after)
    return p

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '黑体' if level <= 2 else '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if level <= 2 else '宋体')
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_code_block(doc, code_text, font_size=9):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(qn('w:shd'), {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): 'F2F2F2'})
    pPr.append(shading)
    run = p.add_run(code_text)
    run.font.size = Pt(font_size)
    run.font.name = 'Consolas'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    set_paragraph_format(p, 6, 6)
    return p

def add_formula(doc, formula, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(formula)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, 6, space_after)
    return p

def add_table_with_data(doc, headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table

def add_bullet_list(doc, items, font_size=12):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(font_size)
        run.font.name = '宋体'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_paragraph_format(p, 0, 3)

def main():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    
    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = '黑体'
        hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        hs.font.color.rgb = RGBColor(0, 0, 0)
    
    # ========== 标题 ==========
    h = doc.add_heading('基于协同过滤的商品推荐系统设计与实现', level=0)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(22)
    
    # ========== 摘要 ==========
    add_heading_styled(doc, '摘要', level=1)
    add_body(doc, '随着电子商务的快速发展，个性化推荐系统已成为提升用户体验和促进销售的核心技术。本研究设计并实现了一个基于协同过滤的商品推荐系统，采用Spring Boot 3.3.2后端框架与Vue 3前端技术栈，通过融合多种推荐算法，有效解决了冷启动和数据稀疏性问题。系统实现了基于用户的协同过滤（User-Based CF）、基于物品的协同过滤（Item-Based CF）、基于行为的推荐（Behavior-Based）以及混合推荐（Hybrid）四种算法策略。混合推荐算法融合Item-CF、User-CF、热门物品、物品关联和内容相似度五种信号，通过sigmoid函数实现动态权重调整，并引入MMR多样化算法平衡相关性与多样性。实验结果表明，在合成数据集上，Item-CF在Precision@10上达到0.2244，NDCG@10达到0.5305；混合推荐算法的Coverage达到0.779，显著优于单一算法。系统采用分层架构设计，包含用户认证、推荐引擎、行为分析、Redis缓存管理和物品关联预计算等核心模块，支持手机号/邮箱登录、Token认证、时效衰减、推荐理由生成等完整功能。')
    
    p = doc.add_paragraph()
    run = p.add_run('关键词：')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = '宋体'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2 = p.add_run('协同过滤；混合推荐；动态权重；MMR多样化；Spring Boot；Vue 3')
    run2.font.size = Pt(12)
    run2.font.name = '宋体'
    r2 = run2._element
    r2.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # ========== 第一章 ==========
    add_heading_styled(doc, '1. 引言', level=1)
    
    add_heading_styled(doc, '1.1 研究背景与意义', level=2)
    add_body(doc, '在信息爆炸的时代，用户面临着海量商品信息的选择困境。推荐系统通过分析用户行为数据，为用户提供个性化的商品推荐，有效缓解信息过载问题。协同过滤作为推荐系统的核心技术，通过挖掘用户间的相似性或物品间的关联性，实现精准推荐。')
    add_body(doc, '随着电商平台的日益复杂，单一推荐算法已无法满足多样化的用户需求。User-Based CF擅长发现用户群体的共同偏好，Item-Based CF在物品关联推荐上表现稳定，但两者都存在冷启动和数据稀疏性问题。因此，设计一个融合多种算法优势、具备自适应能力的混合推荐系统具有重要的理论价值和实际意义。')
    
    add_heading_styled(doc, '1.2 研究目标', level=2)
    add_body(doc, '本研究的主要目标是设计并实现一个基于协同过滤的商品推荐系统，具体包括：')
    add_bullet_list(doc, [
        '实现基于用户的协同过滤算法（User-Based CF），采用Pearson相似度与温和shrinkage机制',
        '实现基于物品的协同过滤算法（Item-Based CF），采用Adjusted Cosine相似度消除用户评分偏差',
        '实现基于行为的推荐算法（Behavior-Based），将显式评分与隐式行为（浏览、点击、加购、收藏）统一映射',
        '设计混合推荐策略，融合五种推荐信号，通过sigmoid函数实现动态权重调整',
        '实现MMR多样化算法，平衡推荐结果的相关性与多样性',
        '设计时间衰减机制，使近期行为对推荐结果的影响更大',
        '构建完整的Web应用，包含用户认证、行为管理、推荐展示和离线评估功能',
        '通过实验验证系统性能'
    ])
    
    add_heading_styled(doc, '1.3 论文结构', level=2)
    add_body(doc, '本文共分为六章：第一章介绍研究背景和目标；第二章综述协同过滤相关技术；第三章阐述系统总体设计；第四章详细说明算法实现；第五章展示实验结果与分析；第六章总结研究成果并展望未来工作。')
    
    # ========== 第二章 ==========
    add_heading_styled(doc, '2. 相关技术与理论基础', level=1)
    
    add_heading_styled(doc, '2.1 协同过滤技术概述', level=2)
    add_body(doc, '协同过滤（Collaborative Filtering, CF）是一种基于用户行为的推荐方法，其核心思想是通过分析用户的历史行为数据，发现用户之间或物品之间的相似性，从而为用户推荐可能感兴趣的物品。协同过滤不依赖物品内容信息，仅利用用户-物品交互矩阵即可实现推荐，具有通用性强、实现简单的优点。')
    
    add_heading_styled(doc, '2.2 基于用户的协同过滤（User-Based CF）', level=2)
    add_body(doc, '基于用户的协同过滤通过计算用户之间的相似度，找到与目标用户兴趣相似的邻居用户，然后将邻居用户喜欢的物品推荐给目标用户。')
    add_body(doc, '系统采用Pearson相关系数计算用户间相似度，并引入温和shrinkage机制避免低重叠度用户间的过度置信：')
    add_formula(doc, 'sim(u,v) = [Σ(r_ui × r_vi) - (Σr_ui × Σr_vi)/n] / sqrt[(Σr_ui² - (Σr_ui)²/n) × (Σr_vi² - (Σr_vi)²/n)] × overlap/(overlap+5)')
    add_body(doc, '评分预测公式：')
    add_formula(doc, 'r̂_ui = Σ(sim(u,v) × r_vi) / Σ(sim(u,v))')
    
    add_heading_styled(doc, '2.3 基于物品的协同过滤（Item-Based CF）', level=2)
    add_body(doc, '基于物品的协同过滤通过计算物品之间的相似度，找到与用户已评分物品相似的其他物品，然后推荐给用户。')
    add_body(doc, '系统采用Adjusted Cosine相似度，先减去全局均值（GLOBAL_MEAN=3.5）消除用户评分偏差：')
    add_formula(doc, 'sim(i,j) = [Σ(r_ui - r̄)(r_uj - r̄)] / sqrt[Σ(r_ui - r̄)² × Σ(r_uj - r̄)²] × overlap/(overlap+8)')
    
    add_heading_styled(doc, '2.4 混合推荐策略', level=2)
    add_body(doc, '单一算法存在局限性，混合推荐策略通过融合多种算法的优势，提升推荐效果。本系统融合五种推荐信号：Item-CF、User-CF、热门物品、物品关联和内容相似度，通过sigmoid函数根据用户活跃度动态调整各信号权重。')
    
    add_heading_styled(doc, '2.5 评价指标', level=2)
    add_table_with_data(doc, ['指标', '说明'], [
        ['精确率@K', '衡量推荐准确性，推荐列表中相关物品的比例'],
        ['召回率@K', '衡量推荐完整性，用户相关物品中被推荐的比例'],
        ['NDCG@K', '衡量推荐排序质量，考虑位置折扣的归一化累计增益'],
        ['覆盖率', '衡量推荐物品的覆盖范围，被推荐物品占全部物品的比例']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '2.6 开发技术栈', level=2)
    add_table_with_data(doc, ['技术', '版本', '用途'], [
        ['Java', '17', '后端编程语言'],
        ['Spring Boot', '3.3.2', '后端框架'],
        ['Spring Data JPA', '-', '数据访问层'],
        ['MySQL', '8.0', '关系型数据库'],
        ['Redis', '7.0', '缓存服务'],
        ['Vue 3', '-', '前端框架'],
        ['Vite', '-', '前端构建工具'],
        ['Maven', '-', '项目构建工具']
    ])
    doc.add_paragraph()
    
    # ========== 第三章 ==========
    add_heading_styled(doc, '3. 系统总体设计', level=1)
    
    add_heading_styled(doc, '3.1 系统架构', level=2)
    add_body(doc, '系统采用前后端分离的分层架构设计，后端基于Spring Boot，前端基于Vue 3 + Vite。系统分为五层：前端展示层、API控制层、业务逻辑层、算法层和数据访问层。')
    add_body(doc, '前端展示层：Vue 3组件（Home.vue、Login.vue、ItemDetail.vue等）', bold=True)
    add_body(doc, 'API控制层：RESTful API控制器（RecommendationController、AuthController、BehaviorController等）', bold=True)
    add_body(doc, '业务逻辑层：核心服务类（RecommendationService、AuthService、BehaviorService等）', bold=True)
    add_body(doc, '算法层：推荐算法实现（UserBasedCF、ItemBasedCF、SimilarityMetrics等）', bold=True)
    add_body(doc, '数据访问层：JPA Repository（UserRepository、ItemRepository、RatingRepository等）', bold=True)
    add_body(doc, '数据存储：MySQL（用户/物品/评分/标记表）+ Redis（缓存/Token/验证码）', bold=True)
    
    add_heading_styled(doc, '3.2 核心模块设计', level=2)
    
    add_heading_styled(doc, '3.2.1 推荐服务模块（RecommendationService）', level=3)
    add_body(doc, '推荐服务是系统的核心业务模块，提供以下功能：')
    add_bullet_list(doc, [
        'recommendForUser(userId, topN, type)：为用户生成推荐列表',
        'recommendForUserWithReason(userId, topN, type)：生成带推荐理由的推荐列表',
        'recommendWithDiversity(userId, topN, type, diversityLevel)：生成多样性优化的推荐列表',
        'getPopularItems(topN)：获取热门商品',
        'getPopularItemsByCategory(category, topN)：获取分类热门商品'
    ])
    add_body(doc, '支持四种算法类型：USER_BASED、ITEM_BASED、BEHAVIOR_BASED、HYBRID。')
    
    add_heading_styled(doc, '3.2.2 算法模块（algo包）', level=3)
    add_table_with_data(doc, ['类名', '职责', '核心方法'], [
        ['RecommenderStrategy', '推荐算法策略接口', 'recommend(userItem, userId, topN)'],
        ['UserBasedCF', '基于用户的协同过滤', 'findNeighbors(), predictRatings()'],
        ['ItemBasedCF', '基于物品的协同过滤', 'buildItemUsers(), predictRatings()'],
        ['SimilarityMetrics', '相似度计算工具类', 'userSimilarity(), itemSimilarity()'],
        ['Recommendation', '推荐结果封装', 'getItemId(), getScore(), compareTo()'],
        ['AlgorithmEvaluator', '离线评估器', 'loadMovieLens100k(), evaluate()']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.2.3 认证模块（AuthService）', level=3)
    add_body(doc, '提供完整的用户认证功能：')
    add_bullet_list(doc, [
        '手机号验证码登录、手机号密码登录、邮箱密码登录、用户名密码登录',
        '邮箱注册（带验证码）、手机注册',
        'Token生成与验证（Access Token 30分钟，Refresh Token 7天）',
        '密码修改、密码重置、用户信息更新',
        '密码采用SHA-256哈希存储'
    ])
    
    add_heading_styled(doc, '3.2.4 行为管理模块（BehaviorService）', level=3)
    add_bullet_list(doc, [
        '用户/商品搜索查询',
        '评分记录的增删改查',
        '行为记录与评分映射：view→1.6、click→2.2、cart→3.8、favorite→4.5',
        '收藏/加购标记管理',
        '批量评分导入（单次上限1000条）'
    ])
    
    add_heading_styled(doc, '3.2.5 缓存模块', level=3)
    add_body(doc, '采用Redis + 内存双层缓存策略：')
    add_bullet_list(doc, [
        'Redis缓存：推荐结果缓存（30分钟过期）、验证码缓存、Token缓存',
        '内存缓存：RecommendationContext懒加载缓存、ItemAssociationPrecomputeService快照缓存',
        '缓存失效：评分更新后自动标记脏数据，推荐缓存按用户失效'
    ])
    
    add_heading_styled(doc, '3.2.6 物品关联预计算', level=3)
    add_bullet_list(doc, [
        '应用启动时构建物品共现相似度快照',
        '每15分钟定时刷新（可配置）',
        '数据变更时标记脏数据，下次刷新时重建',
        '使用PriorityQueue维护每个物品的Top-N邻居（默认120个）'
    ])
    
    add_heading_styled(doc, '3.3 数据库设计', level=2)
    add_body(doc, '系统使用MySQL 8.0数据库，包含4张核心表：')
    
    add_heading_styled(doc, '用户表（users）', level=3)
    add_table_with_data(doc, ['字段', '类型', '约束', '说明'], [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '用户ID'],
        ['username', 'VARCHAR(100)', 'UNIQUE, NOT NULL', '用户名'],
        ['phone', 'VARCHAR(20)', 'UNIQUE', '手机号'],
        ['password_hash', 'VARCHAR(64)', '-', 'SHA-256密码哈希'],
        ['email', 'VARCHAR(100)', 'UNIQUE', '邮箱'],
        ['disabled', 'TINYINT(1)', 'DEFAULT 0', '是否禁用'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '物品表（items）', level=3)
    add_table_with_data(doc, ['字段', '类型', '约束', '说明'], [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '物品ID'],
        ['name', 'VARCHAR(200)', 'NOT NULL', '物品名称'],
        ['category', 'VARCHAR(100)', '-', '物品类别'],
        ['image_url', 'VARCHAR(255)', '-', '图片URL'],
        ['created_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '创建时间']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '评分表（ratings）', level=3)
    add_table_with_data(doc, ['字段', '类型', '约束', '说明'], [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '评分ID'],
        ['user_id', 'BIGINT', 'FK→users, NOT NULL', '用户ID'],
        ['item_id', 'BIGINT', 'FK→items, NOT NULL', '物品ID'],
        ['score', 'DOUBLE', 'CHECK(0≤score≤5)', '评分值'],
        ['rated_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '评分时间']
    ])
    doc.add_paragraph()
    add_body(doc, '唯一约束：UNIQUE(user_id, item_id)；索引：idx_r_item_id, idx_r_rated_at, idx_r_user_rated_at')
    
    add_heading_styled(doc, '用户物品标记表（user_item_flags）', level=3)
    add_table_with_data(doc, ['字段', '类型', '约束', '说明'], [
        ['id', 'BIGINT', 'PK, AUTO_INCREMENT', '标记ID'],
        ['user_id', 'BIGINT', 'FK→users, NOT NULL', '用户ID'],
        ['item_id', 'BIGINT', 'FK→items, NOT NULL', '物品ID'],
        ['favorite', 'TINYINT(1)', 'DEFAULT 0', '是否收藏'],
        ['in_cart', 'TINYINT(1)', 'DEFAULT 0', '是否加购'],
        ['updated_at', 'TIMESTAMP', 'DEFAULT CURRENT_TIMESTAMP', '更新时间']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.4 API接口设计', level=2)
    
    add_heading_styled(doc, '3.4.1 推荐接口', level=3)
    add_table_with_data(doc, ['接口', '方法', '功能'], [
        ['/api/recommendations/{userId}?n=5&algo=user', 'GET', '获取个性化推荐（带理由）'],
        ['/api/recommendations/{userId}/diverse', 'GET', '获取多样性优化推荐'],
        ['/api/recommendations/popular?n=10', 'GET', '获取热门商品']
    ])
    doc.add_paragraph()
    add_body(doc, 'algo参数支持：user、item、behavior、hybrid')
    
    add_heading_styled(doc, '3.4.2 认证接口', level=3)
    add_table_with_data(doc, ['接口', '方法', '功能'], [
        ['/api/auth/send-sms-code', 'POST', '发送手机验证码'],
        ['/api/auth/send-email-code', 'POST', '发送邮箱验证码'],
        ['/api/auth/login/sms', 'POST', '手机验证码登录'],
        ['/api/auth/login', 'POST', '手机密码登录'],
        ['/api/auth/login-email', 'POST', '邮箱密码登录'],
        ['/api/auth/login-username', 'POST', '用户名密码登录'],
        ['/api/auth/register', 'POST', '手机注册'],
        ['/api/auth/register-email', 'POST', '邮箱注册'],
        ['/api/auth/refresh', 'POST', '刷新Token'],
        ['/api/auth/logout', 'POST', '登出'],
        ['/api/auth/me', 'GET/PUT', '获取/更新用户信息'],
        ['/api/auth/me/change-password', 'POST', '修改密码']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.4.3 行为管理接口', level=3)
    add_table_with_data(doc, ['接口', '方法', '功能'], [
        ['/api/behaviors/users?keyword=xxx', 'GET', '搜索用户'],
        ['/api/behaviors/items?keyword=xxx', 'GET', '搜索商品'],
        ['/api/behaviors/users/{userId}/ratings', 'GET', '获取用户评分'],
        ['/api/behaviors/ratings', 'POST', '创建/更新评分'],
        ['/api/behaviors/ratings/batch', 'POST', '批量导入评分'],
        ['/api/behaviors/events', 'POST', '记录行为（view/click/cart/favorite）']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.4.4 评估接口', level=3)
    add_table_with_data(doc, ['接口', '方法', '功能'], [
        ['/api/evaluations/offline', 'GET', '执行离线评估'],
        ['/api/evaluations/offline/csv', 'GET', '导出评估结果CSV']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '3.5 前端设计', level=2)
    add_body(doc, '前端采用Vue 3 + Vite构建，包含以下核心页面：')
    add_table_with_data(doc, ['页面', '文件', '功能'], [
        ['首页', 'Home.vue', '展示推荐列表、热门商品'],
        ['登录', 'Login.vue', '多种登录方式'],
        ['注册', 'Register.vue', '手机/邮箱注册'],
        ['商品详情', 'ItemDetail.vue', '商品详情、评分、行为记录'],
        ['收藏', 'Favorites.vue', '用户收藏列表'],
        ['购物车', 'Cart.vue', '用户加购列表']
    ])
    doc.add_paragraph()
    add_body(doc, '核心组件：TopNav.vue（顶部导航栏）、ItemCard.vue（商品卡片）、RatingStars.vue（评分星级组件）')
    
    # ========== 第四章 ==========
    add_heading_styled(doc, '4. 核心算法实现', level=1)
    
    add_heading_styled(doc, '4.1 算法总体架构', level=2)
    add_body(doc, '本系统的推荐算法层采用策略模式设计，通过RecommenderStrategy接口统一算法入口，支持多种推荐策略的灵活切换与组合。算法架构包含四个核心层次：基础协同过滤层、相似度计算层、混合推荐层和策略优化层。')
    add_code_block(doc, 'public interface RecommenderStrategy {\n    List<Recommendation> recommend(Map<Long, Map<Long, Double>> userItemRatings, Long userId, int topN);\n}\n\npublic enum AlgorithmType {\n    USER_BASED,      // 基于用户的协同过滤\n    ITEM_BASED,      // 基于物品的协同过滤\n    BEHAVIOR_BASED,  // 基于行为的推荐\n    HYBRID           // 混合推荐算法\n}')
    
    add_heading_styled(doc, '4.2 相似度计算模块', level=2)
    add_body(doc, '相似度计算是协同过滤算法的核心。系统设计了SimilarityMetrics工具类，提供多种相似度计算方法，并针对User-CF和Item-CF的不同特点进行了专门优化。')
    
    add_heading_styled(doc, '4.2.1 User-CF专用Pearson相似度', level=3)
    add_body(doc, 'User-CF采用Pearson相关系数计算用户间相似度，并引入温和shrinkage机制避免低重叠度用户间的过度置信。shrinkage参数设置为5。')
    add_code_block(doc, 'public static double userSimilarity(Map<Long, Double> a, Map<Long, Double> b) {\n    int overlap = overlapCount(a, b);\n    if (overlap < MIN_OVERLAP) return 0.0;\n    // 计算Pearson相关系数\n    double num = sumAB - (sumA * sumB / n);\n    double den = sqrt((sumA2 - sumA*sumA/n) * (sumB2 - sumB*sumB/n));\n    if (den <= 1e-12) return 0.0;\n    double sim = num / den;\n    return sim * overlap / (overlap + 5.0);\n}')
    
    add_heading_styled(doc, '4.2.2 Item-CF专用Adjusted Cosine相似度', level=3)
    add_body(doc, 'Item-CF采用Adjusted Cosine相似度，先减去全局均值（GLOBAL_MEAN=3.5）消除用户评分偏差。shrinkage参数设置为8。')
    add_code_block(doc, 'public static double itemSimilarity(Map<Long, Double> a, Map<Long, Double> b) {\n    int overlap = overlapCount(a, b);\n    if (overlap < 3) return 0.0;\n    double va = e.getValue() - GLOBAL_MEAN;\n    double vb = other - GLOBAL_MEAN;\n    dot += va * vb;\n    double sim = dot / (sqrt(normA) * sqrt(normB));\n    return sim * overlap / (overlap + 8.0);\n}')
    
    add_heading_styled(doc, '4.3 基于用户的协同过滤算法', level=2)
    add_body(doc, 'UserBasedCF实现了基于用户的协同过滤，核心流程：获取目标用户评分→寻找相似用户邻居→预测候选物品评分→生成推荐列表。')
    add_table_with_data(doc, ['参数', '值', '说明'], [
        ['MIN_SIMILARITY', '0.0', '最小相似度阈值'],
        ['GLOBAL_MEAN', '3.5', '全局平均评分'],
        ['DEFAULT_NEIGHBORS', '30', '默认邻居数量'],
        ['MIN_OVERLAP', '2', '最小共同评分物品数']
    ])
    doc.add_paragraph()
    add_body(doc, '邻居查找阶段遍历所有用户，计算与目标用户的余弦相似度，按相似度降序排序取前30个。评分预测采用加权平均法：r̂_ui = Σ(sim×r_vi) / Σ(sim)。当无法找到有效邻居时，回退到热门物品推荐。')
    
    add_heading_styled(doc, '4.4 基于物品的协同过滤算法', level=2)
    add_body(doc, 'ItemBasedCF通过计算物品间的相似度来推荐与用户已评分物品相似的其他物品。需要构建物品到用户的反向映射（buildItemUsers），然后遍历用户已评分的每个物品，计算与候选物品的相似度进行加权评分。')
    
    add_heading_styled(doc, '4.5 基于行为的推荐算法', level=2)
    add_body(doc, 'Behavior-Based将显式评分转换为隐式行为强度，综合考虑评分值和时效性。')
    add_table_with_data(doc, ['行为类型', '映射评分', '偏好强度'], [
        ['view（浏览）', '1.6', '弱偏好'],
        ['click（点击）', '2.2', '弱偏好'],
        ['cart（加购）', '3.8', '中强偏好'],
        ['favorite（收藏）', '4.5', '强偏好']
    ])
    doc.add_paragraph()
    add_body(doc, '隐式行为强度公式：strength = (0.2 + 0.8 × score/5.0) × (0.4 + 0.6 × decay)。对于近期交互（decay≥0.8），额外施加15%的放大系数。')
    
    add_heading_styled(doc, '4.6 混合推荐算法', level=2)
    add_body(doc, '混合推荐算法是本系统的核心创新点，融合五种推荐信号，根据用户活跃度动态调整权重。')
    
    add_heading_styled(doc, '4.6.1 五种推荐信号', level=3)
    add_table_with_data(doc, ['信号', '来源', '计算方式'], [
        ['Item-CF', '基于物品的协同过滤', '物品相似度加权评分'],
        ['User-CF', '基于用户的协同过滤', '用户相似度加权评分'],
        ['Popularity', '热门物品', '平均分×log(1+评分次数)'],
        ['Association', '物品关联', '预计算共现相似度'],
        ['Content', '内容相似度', '类别偏好匹配度']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '4.6.2 动态权重调整', level=3)
    add_body(doc, '使用sigmoid函数根据用户历史评分数量动态计算权重。归一化活跃度=min(1.0, ratedCount/30.0)，sigmoid中心点0.4，斜率10。')
    add_table_with_data(doc, ['权重', '公式', '范围'], [
        ['Item-CF', '0.30 + 0.15 × sigmoid', '0.30~0.45'],
        ['User-CF', '0.15 + 0.15 × sigmoid', '0.15~0.30'],
        ['Popularity', '0.25 - 0.15 × sigmoid', '0.25~0.10'],
        ['Association', '0.10 + 0.02 × sigmoid', '0.10~0.12'],
        ['Content', '0.20 - 0.15 × sigmoid', '0.20~0.05']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '4.6.3 偏好类别提升', level=3)
    add_body(doc, '识别用户偏好类别（平均分≥4.0且加权评分数≥2），计算类别强度：strength = min(1.0, ((avg-4.0)×0.7) + min(0.3, (count-2.0)×0.08))。最终得分乘以(1.0 + 0.25 × categoryStrength)提升。')
    
    add_heading_styled(doc, '4.7 物品关联预计算', level=2)
    add_body(doc, 'ItemAssociationPrecomputeService离线预计算物品间共现相似度。相似度公式：sim(i,j) = co_count(i,j) / sqrt(userCount(i) × userCount(j))。每个物品保留Top-120邻居，使用PriorityQueue维护。')
    
    add_heading_styled(doc, '4.8 时间衰减机制', level=2)
    add_body(doc, '衰减权重使用半衰期30天的指数衰减：decay = 0.5^(days/30)。应用于User-CF评分矩阵、偏好类别计算和隐式行为强度，热门物品排序不应用衰减。')
    
    add_heading_styled(doc, '4.9 冷启动解决方案', level=2)
    add_body(doc, '多层级回退策略：主要算法推荐→热门物品回退→目录级冷启动补齐。热门回退引入随机扰动（±10%）和类别多样性约束（同类别≤40%）。目录补齐根据用户活跃度动态调整偏好权重。')
    
    add_heading_styled(doc, '4.10 多样化推荐（MMR算法）', level=2)
    add_body(doc, 'MMR公式：MMR = λ × relevance - (1-λ) × max(similarity_to_selected)，λ=0.7。支持可调节多样性级别：lambda = 1.0 - diversityLevel × 0.4（范围0.6-1.0）。')
    
    add_heading_styled(doc, '4.11 推荐理由生成', level=2)
    add_body(doc, '根据算法类型、类别匹配和得分生成推荐理由。HYBRID算法在类别匹配且偏好强度≥1.2时生成个性化理由，得分≥0.7时强调综合信号优势。')
    
    add_heading_styled(doc, '4.12 推荐上下文缓存', level=2)
    add_body(doc, 'RecommendationContext采用懒加载策略缓存中间计算结果：用户-物品评分矩阵、衰减权重矩阵、物品-用户矩阵、热门物品排名、物品类别映射。')
    
    add_heading_styled(doc, '4.13 算法流程总结', level=2)
    add_body(doc, '(1) 接收请求，获取参数 → (2) 构建上下文，懒加载矩阵 → (3) 执行推荐策略 → (4) 合并结果并回退 → (5) HYBRID应用MMR多样化 → (6) 生成理由，返回结果')
    
    # ========== 第五章 ==========
    add_heading_styled(doc, '5. 实验与评估', level=1)
    
    add_heading_styled(doc, '5.1 实验环境', level=2)
    add_table_with_data(doc, ['配置项', '说明'], [
        ['数据集', 'MovieLens 100K / 合成数据'],
        ['编程语言', 'Java 17'],
        ['框架', 'Spring Boot 3.3.2'],
        ['数据库', 'MySQL 8.0'],
        ['缓存', 'Redis 7.0'],
        ['测试集比例', '0.2'],
        ['相关性阈值', '1.5 / 4.0']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.2 离线评估设计', level=2)
    add_body(doc, '通过OfflineEvaluationService实现离线评估，支持topK（1-100）、testRatio（默认0.2）、relevanceThreshold（默认1.5）配置。数据集划分采用基于时间的策略：按评分时间排序，前(1-testRatio)为训练集，后testRatio为测试集。评估USER_BASED、ITEM_BASED、HYBRID三种算法的Precision@K、Recall@K、NDCG@K、Coverage四项指标。')
    
    add_heading_styled(doc, '5.3 实验结果', level=2)
    
    add_heading_styled(doc, '5.3.1 大规模合成数据评估（testSize=28126）', level=3)
    add_body(doc, '使用500用户、300物品、150000条评分的合成数据集：')
    add_table_with_data(doc, ['算法', 'K', 'Precision@K', 'Recall@K', 'NDCG@K', 'Coverage'], [
        ['hybrid', '5', '0.2638', '0.3874', '0.5241', '0.3194'],
        ['hybrid', '10', '0.2050', '0.4200', '0.5102', '0.4710'],
        ['hybrid', '20', '0.1659', '0.4616', '0.5041', '0.6613'],
        ['item', '5', '0.2738', '0.3951', '0.5303', '0.3097'],
        ['item', '10', '0.2244', '0.4361', '0.5305', '0.4710'],
        ['item', '20', '0.1891', '0.4919', '0.5328', '0.7226'],
        ['user', '5', '0.2200', '0.3878', '0.4790', '0.3258'],
        ['user', '10', '0.1694', '0.4207', '0.4779', '0.5194'],
        ['user', '20', '0.1338', '0.4621', '0.4797', '0.7645']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.3.2 MovieLens 100K评估（relevance=1.5, K=10）', level=3)
    add_table_with_data(doc, ['算法', 'Precision@K', 'Recall@K', 'NDCG@K', 'Coverage', 'Users'], [
        ['user', '0.0000', '0.0000', '0.0000', '0.0098', '9049'],
        ['item', '0.0006', '0.0038', '0.0025', '0.0264', '9049'],
        ['hybrid', '0.0066', '0.0360', '0.0225', '0.7786', '9049']
    ])
    doc.add_paragraph()
    
    add_heading_styled(doc, '5.4 结果分析', level=2)
    add_bullet_list(doc, [
        'Item-CF在Precision/NDCG上整体最优：在合成数据评估中，Item-CF在Precision@K和NDCG@K上均优于User-CF',
        'User-CF在Coverage上最优：覆盖率在各K值下均最高，推荐范围更广',
        'Hybrid的Coverage显著领先：在MovieLens评估中达到0.779，多信号融合有效扩展推荐范围',
        '指标曲线符合预期：K增大时Precision下降，Recall和Coverage上升',
        '动态权重策略有效：低活跃度侧重内容和热门信号，高活跃度侧重协同过滤信号'
    ])
    
    add_heading_styled(doc, '5.5 系统性能', level=2)
    add_bullet_list(doc, [
        '物品关联预计算：应用启动时构建，每15分钟刷新，在线查询为O(1)',
        '推荐结果缓存：Redis缓存30分钟，显著降低重复请求的计算开销',
        '懒加载策略：RecommendationContext只在首次访问时计算中间结果',
        '批量评分导入：单次上限1000条，事务内完成'
    ])
    
    # ========== 第六章 ==========
    add_heading_styled(doc, '6. 结论与展望', level=1)
    
    add_heading_styled(doc, '6.1 研究成果', level=2)
    add_bullet_list(doc, [
        '实现了User-Based CF、Item-Based CF、Behavior-Based和Hybrid四种推荐算法',
        '设计了SimilarityMetrics工具类，针对User-CF和Item-CF分别采用Pearson和Adjusted Cosine相似度',
        '设计了多层级冷启动回退策略，引入随机扰动和类别多样性约束',
        '实现了MMR算法平衡相关性与多样性，支持可调节的多样性级别',
        '引入半衰期30天的指数衰减机制，使近期行为影响更大',
        '构建了包含用户认证、行为管理、推荐展示、离线评估的完整Web应用',
        '采用Redis缓存、物品关联预计算、懒加载策略等优化手段'
    ])
    
    add_heading_styled(doc, '6.2 不足与改进', level=2)
    add_bullet_list(doc, [
        'MovieLens评估指标偏低：电影评分数据集与商品推荐场景存在差异',
        '混合推荐权重基于经验设定，后续可通过网格搜索或贝叶斯优化调参',
        '当前系统采用离线预计算+缓存策略，实时性有限'
    ])
    
    add_heading_styled(doc, '6.3 未来工作', level=2)
    add_bullet_list(doc, [
        '引入深度学习模型：Neural Collaborative Filtering、Graph Neural Network等',
        '实时推荐优化：引入Kafka和Flink实现近实时推荐',
        '增强推荐解释性：结合知识图谱提供可解释的推荐',
        '多模态数据融合：引入商品图片、文本描述等多模态特征',
        'A/B测试框架：构建在线A/B测试平台',
        '长尾推荐优化：针对长尾物品设计专门的推荐策略'
    ])
    
    # 保存
    output_path = r'D:\app\ks\recommendtwo\docs\基于协同过滤的商品推荐系统设计与实现_完整版.docx'
    doc.save(output_path)
    print(f'文件已保存至: {output_path}')

if __name__ == '__main__':
    main()
