#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import sys
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
import re

def md_to_docx(md_file_path, docx_file_path):
    """Convert Markdown file to DOCX format with proper formatting"""

    # Read markdown content
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Create document
    doc = Document()

    # Set page setup - A4 paper
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Title style (Level 1)
    title_style = doc.styles['Heading 1']
    title_font = title_style.font
    title_font.name = '黑体'
    title_font.size = Pt(18)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 0, 0)
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.line_spacing = 1.5
    title_style.paragraph_format.space_before = Pt(12)
    title_style.paragraph_format.space_after = Pt(12)

    # Heading 2 style
    h2_style = doc.styles['Heading 2']
    h2_font = h2_style.font
    h2_font.name = '黑体'
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 0, 0)
    h2_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2_style.paragraph_format.line_spacing = 1.5
    h2_style.paragraph_format.space_before = Pt(10)
    h2_style.paragraph_format.space_after = Pt(10)

    # Heading 3 style
    h3_style = doc.styles['Heading 3']
    h3_font = h3_style.font
    h3_font.name = '黑体'
    h3_font.size = Pt(12)
    h3_font.bold = True
    h3_font.color.rgb = RGBColor(0, 0, 0)
    h3_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3_style.paragraph_format.line_spacing = 1.5
    h3_style.paragraph_format.space_before = Pt(8)
    h3_style.paragraph_format.space_after = Pt(8)

    # Heading 4 style
    h4_style = doc.styles['Heading 4']
    h4_font = h4_style.font
    h4_font.name = '黑体'
    h4_font.size = Pt(11)
    h4_font.bold = True
    h4_font.color.rgb = RGBColor(0, 0, 0)
    h4_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h4_style.paragraph_format.line_spacing = 1.5
    h4_style.paragraph_format.space_before = Pt(6)
    h4_style.paragraph_format.space_after = Pt(6)

    # List styles
    list_style = doc.styles['List Bullet']
    list_font = list_style.font
    list_font.name = '宋体'
    list_font.size = Pt(12)
    list_style.paragraph_format.line_spacing = 1.5
    list_style.paragraph_format.first_line_indent = Cm(0)

    num_list_style = doc.styles['List Number']
    num_list_font = num_list_style.font
    num_list_font.name = '宋体'
    num_list_font.size = Pt(12)
    num_list_style.paragraph_format.line_spacing = 1.5
    num_list_style.paragraph_format.first_line_indent = Cm(0)

    # Process content line by line from original markdown
    lines = md_content.split('\n')

    in_code_block = False
    code_content = []
    in_math_block = False
    math_content = []

    # Track chapter numbering
    chapter_num = 0
    section_num = 0
    subsection_num = 0

    for line in lines:
        line = line.rstrip()

        # Skip empty lines but add space
        if not line or line.strip() == '':
            if not in_code_block and not in_math_block:
                doc.add_paragraph()
            continue

        # Code block detection
        if line.startswith('```'):
            in_code_block = not in_code_block
            if in_code_block:
                code_content = []
            else:
                # Add code block
                p = doc.add_paragraph()
                p.style = 'Normal'
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Add code with monospace font
                for code_line in code_content:
                    run = p.add_run(code_line + '\n')
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Math block detection
        if line.startswith('$$'):
            in_math_block = not in_math_block
            if in_math_block:
                math_content = []
            else:
                math_text = ' '.join(math_content)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(12)
                run = p.add_run(math_text)
                run.font.name = 'Cambria Math'
                run.font.size = Pt(11)
                run.italic = True
            continue

        if in_math_block:
            math_content.append(line)
            continue

        # Horizontal rule (---)
        if line.startswith('---') and len(line) >= 3:
            # Add space instead of horizontal rule
            doc.add_paragraph()
            continue

        # Headings
        if line.startswith('#### '):
            heading = line[5:].strip()
            p = doc.add_heading(heading, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if line.startswith('### '):
            heading = line[4:].strip()
            p = doc.add_heading(heading, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            continue

        if line.startswith('## '):
            heading = line[3:].strip()
            # Check for numbered headings like "1. 引言"
            match = re.match(r'^(\d+)\.\s+(.+)', heading)
            if match:
                chapter_num = int(match.group(1))
                heading = match.group(2)
                section_num = 0
            p = doc.add_heading(f"{chapter_num}. {heading}", level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        if line.startswith('# '):
            heading = line[2:].strip()
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue

        # Tables
        if line.startswith('|'):
            cells = [c.strip() for c in line.split('|')[1:-1]]
            if len(cells) > 0:
                # Check if this is a separator line (contains ---)
                is_separator = any('---' in cell for cell in cells)
                if is_separator:
                    continue

                # Check if we need to create a new table or add to existing
                last_table = doc.tables[-1] if doc.tables else None
                is_header = False

                if last_table:
                    try:
                        if len(last_table.columns) == len(cells):
                            row_cells = last_table.add_row().cells
                            for i, cell in enumerate(cells):
                                row_cells[i].text = cell
                                row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            continue
                    except:
                        pass

                # Create new table
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells = table.rows[0].cells
                for i, cell in enumerate(cells):
                    row_cells[i].text = cell
                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    # Make header bold
                    for paragraph in row_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.name = '宋体'
                            run.font.size = Pt(11)
                # Set table font for remaining cells
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = '宋体'
                                run.font.size = Pt(11)
            continue

        # Lists (bullet points)
        if line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.left_indent = Cm(0.5)
            # Handle bold text within list items
            if '**' in text:
                parts = text.split('**')
                for i, part in enumerate(parts):
                    if i % 2 == 1:
                        run = p.add_run(part)
                        run.bold = True
                    else:
                        p.add_run(part)
            else:
                p.add_run(text)
            continue

        # Numbered lists
        if line[0].isdigit() and line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.')):
            parts = line.split('.', 1)
            if len(parts) > 1:
                text = parts[1].strip()
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.left_indent = Cm(0.5)
                # Handle bold text within list items
                if '**' in text:
                    parts_text = text.split('**')
                    for i, part in enumerate(parts_text):
                        if i % 2 == 1:
                            run = p.add_run(part)
                            run.bold = True
                        else:
                            p.add_run(part)
                else:
                    p.add_run(text)
            else:
                p = doc.add_paragraph(line)
            continue

        # Bold text
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 1:
                    run = p.add_run(part)
                    run.bold = True
                else:
                    p.add_run(part)
            continue

        # Math formulas (LaTeX inline)
        if line.startswith('$') and line.endswith('$'):
            p = doc.add_paragraph(line[1:-1])
            continue

        # Regular paragraph
        p = doc.add_paragraph(line)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Cm(0.74)

    # Add table of contents
    # First, we need to add TOC at the beginning
    # Let's create a new document and insert TOC
    # This is complex, we'll add a placeholder TOC section

    # Save document
    doc.save(docx_file_path)
    print(f"Successfully converted {md_file_path} to {docx_file_path}")

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_file = sys.argv[1]
    docx_file = sys.argv[2]

    if not os.path.exists(md_file):
        print(f"Error: Input file '{md_file}' not found")
        sys.exit(1)

    md_to_docx(md_file, docx_file)
