#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成更新后的第四章文档
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading_elm = cell._tc.get_or_add_tcPr()
    shading = shading_elm.makeelement(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd',
        {
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val': 'clear',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color': 'auto',
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill': color
        }
    )
    shading_elm.append(shading)

def add_formatted_paragraph(doc, text, style='Normal', bold=False, italic=False, font_size=12, alignment=None, space_before=0, space_after=6, font_name='宋体', color=None):
    """添加格式化的段落"""
    p = doc.add_paragraph()
    if style != 'Normal':
        p.style = doc.styles[style]
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    run.font.name = font_name
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    return p

def add_code_block(doc, code_text, font_size=9):
    """添加代码块"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    pf.left_indent = Cm(1)
    
    # 设置浅灰色背景
    pPr = p._element.get_or_add_pPr()
    shading = pPr.makeelement(
        qn('w:shd'),
        {
            qn('w:val'): 'clear',
            qn('w:color'): 'auto',
            qn('w:fill'): 'F2F2F2'
        }
    )
    pPr.append(shading)
    
    run = p.add_run(code_text)
    run.font.size = Pt(font_size)
    run.font.name = 'Consolas'
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    return p

def add_table_with_data(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, 'D9E2F3')
    
    # 数据行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置列宽
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Cm(width)
    
    return table

def main():
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.5
    
    # 设置各级标题样式
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = '黑体'
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    # ========== 第四章标题 ==========
    h = doc.add_heading('第四章 核心算法实现', level=1)
    for run in h.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # ========== 4.1 算法总体架构 ==========
    doc.add_heading('4.1 算法总体架构', level=2)
    
    add_formatted_paragraph(doc, 
        '本系统的推荐算法层采用策略模式设计，通过RecommenderStrategy接口统一算法入口，支持多种推荐策略的灵活切换与组合。算法架构包含四个核心层次：基础协同过滤层、相似度计算层、混合推荐层和策略优化层。',
        font_size=12, space_after=6)
    
    add_formatted_paragraph(doc, '算法类型枚举定义如下：', font_size=12, space_after=6)
    
    add_code_block(doc, 
"""public enum AlgorithmType {
    USER_BASED,      // 基于用户的协同过滤
    ITEM_BASED,      // 基于物品的协同过滤
    BEHAVIOR_BASED,  // 基于行为的推荐
    HYBRID           // 混合推荐算法
}""")
    
    add_formatted_paragraph(doc,
        '系统支持四种推荐算法类型，其中USER_BASED和ITEM_BASED为基础协同过滤算法，BEHAVIOR_BASED将显式评分转换为隐式行为强度进行推荐，HYBRID则融合多种信号实现综合推荐。',
        font_size=12, space_after=6)
    
    # ========== 4.2 相似度计算模块 ==========
    doc.add_heading('4.2 相似度计算模块', level=2)
    
    add_formatted_paragraph(doc,
        '相似度计算是协同过滤算法的核心。系统设计了SimilarityMetrics工具类，提供多种相似度计算方法，并针对User-CF和Item-CF的不同特点进行了专门优化。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.2.1 User-CF专用Pearson相似度', level=3)
    
    add_formatted_paragraph(doc,
        'User-CF采用Pearson相关系数计算用户间相似度，并引入温和shrinkage机制避免低重叠度用户间的过度置信。由于用户间评分重叠度通常较低，shrinkage参数设置为5，以平衡置信度与相似度。',
        font_size=12, space_after=6)
    
    add_formatted_paragraph(doc, '计算公式如下：', font_size=12, space_after=6, italic=True)
    
    add_formatted_paragraph(doc,
        'sim(u,v) = [Σ(r_ui × r_vi) - (Σr_ui × Σr_vi)/n] / sqrt[(Σr_ui² - (Σr_ui)²/n) × (Σr_vi² - (Σr_vi)²/n)] × overlap/(overlap+5)',
        font_size=11, space_after=6, italic=True)
    
    add_code_block(doc,
"""public static double userSimilarity(Map<Long, Double> a, Map<Long, Double> b) {
    int overlap = overlapCount(a, b);
    if (overlap < MIN_OVERLAP) return 0.0;  // MIN_OVERLAP = 3
    
    // 计算Pearson相关系数
    double num = sumAB - (sumA * sumB / n);
    double den = sqrt((sumA2 - sumA*sumA/n) * (sumB2 - sumB*sumB/n));
    if (den <= 1e-12) return 0.0;
    
    double sim = num / den;
    // 温和shrinkage：sim * overlap / (overlap + 5)
    return sim * overlap / (overlap + 5.0);
}""")
    
    doc.add_heading('4.2.2 Item-CF专用Adjusted Cosine相似度', level=3)
    
    add_formatted_paragraph(doc,
        'Item-CF采用Adjusted Cosine相似度，先减去全局均值（GLOBAL_MEAN=3.5）消除用户评分偏差，再计算余弦相似度。由于物品间共同用户通常较多，shrinkage参数设置为8。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""public static double itemSimilarity(Map<Long, Double> a, Map<Long, Double> b) {
    int overlap = overlapCount(a, b);
    if (overlap < 3) return 0.0;
    
    // 减去全局均值，消除用户评分偏差
    double va = e.getValue() - GLOBAL_MEAN;  // GLOBAL_MEAN = 3.5
    double vb = other - GLOBAL_MEAN;
    dot += va * vb;
    
    double sim = dot / (sqrt(normA) * sqrt(normB));
    // shrinkage：sim * overlap / (overlap + 8)
    return sim * overlap / (overlap + 8.0);
}""")
    
    doc.add_heading('4.2.3 优化相似度计算', level=3)
    
    add_formatted_paragraph(doc,
        '系统还实现了optimizedSimilarity方法，综合使用Pearson、Adjusted Cosine和标准Cosine三种相似度，当一种方法失效时自动回退到下一种，最后通过置信度加权得到最终相似度。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""public static double optimizedSimilarity(Map<Long, Double> a, Map<Long, Double> b) {
    int overlap = overlapCount(a, b);
    if (overlap < MIN_OVERLAP) return 0.0;
    
    double sim = pearson(a, b);
    if (abs(sim) < 1e-6) sim = adjustedCosine(a, b);
    if (abs(sim) < 1e-6) sim = cosine(a, b);
    
    return confidenceWeighted(sim, overlap);
}

private static double confidenceWeighted(double sim, int overlap) {
    double confidence = log1p(overlap) / log1p(100);
    return sim * confidence;
}""")
    
    # ========== 4.3 基于用户的协同过滤算法 ==========
    doc.add_heading('4.3 基于用户的协同过滤算法（User-Based CF）', level=2)
    
    add_formatted_paragraph(doc,
        'UserBasedCF类实现了基于用户的协同过滤算法，其核心流程包括：获取目标用户评分、寻找相似用户邻居、预测候选物品评分、生成推荐列表。当无法找到足够相似用户时，算法会自动回退到热门物品推荐。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.3.1 核心参数', level=3)
    
    headers = ['参数', '值', '说明']
    rows = [
        ['MIN_SIMILARITY', '0.0', '最小相似度阈值'],
        ['GLOBAL_MEAN', '3.5', '全局平均评分'],
        ['DEFAULT_NEIGHBORS', '30', '默认邻居数量'],
        ['MIN_OVERLAP', '2', '最小共同评分物品数'],
    ]
    add_table_with_data(doc, headers, rows)
    doc.add_paragraph()
    
    doc.add_heading('4.3.2 邻居查找', level=3)
    
    add_formatted_paragraph(doc,
        '邻居查找阶段遍历所有用户，计算与目标用户的余弦相似度。只有共同评分物品数达到MIN_OVERLAP且相似度大于MIN_SIMILARITY的用户才会被纳入候选邻居。候选邻居按相似度降序排序，取前DEFAULT_NEIGHBORS个作为最终邻居。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private List<Neighbor> findNeighbors(Map<Long, Map<Long, Double>> userItem,
                                     Map<Long, Double> targetRatings, Long userId) {
    List<Neighbor> neighbors = new ArrayList<>();
    for (Map.Entry<Long, Map<Long, Double>> entry : userItem.entrySet()) {
        Long otherId = entry.getKey();
        if (Objects.equals(otherId, userId)) continue;
        
        // 计算共同评分物品数
        int overlap = countOverlap(targetRatings, entry.getValue());
        if (overlap < MIN_OVERLAP) continue;
        
        // 计算余弦相似度
        double similarity = cosineSimilarity(targetRatings, entry.getValue());
        if (similarity <= MIN_SIMILARITY) continue;
        
        neighbors.add(new Neighbor(otherId, similarity, entry.getValue()));
    }
    // 按相似度降序排序，取前N个
    neighbors.sort((a, b) -> Double.compare(b.similarity, a.similarity));
    return neighbors.subList(0, min(DEFAULT_NEIGHBORS, neighbors.size()));
}""")
    
    doc.add_heading('4.3.3 评分预测', level=3)
    
    add_formatted_paragraph(doc,
        '评分预测采用加权平均法，以邻居用户的相似度作为权重，对候选物品的评分进行加权平均。预测公式为：',
        font_size=12, space_after=6)
    
    add_formatted_paragraph(doc,
        'pred(i) = Σ(sim_v × r_vi) / Σ(sim_v)，其中sim_v为邻居v的相似度，r_vi为邻居v对物品i的评分',
        font_size=11, space_after=6, italic=True)
    
    add_code_block(doc,
"""private Map<Long, Double> predictRatings(List<Neighbor> neighbors,
                                         Map<Long, Double> targetRatings) {
    Map<Long, Double> predictions = new HashMap<>();
    Map<Long, Double> weightSums = new HashMap<>();
    
    for (Neighbor neighbor : neighbors) {
        double sim = neighbor.similarity;
        if (sim <= 0) continue;
        
        for (Map.Entry<Long, Double> entry : neighbor.ratings.entrySet()) {
            Long itemId = entry.getKey();
            if (targetRatings.containsKey(itemId)) continue;  // 排除已评分物品
            
            predictions.merge(itemId, sim * entry.getValue(), Double::sum);
            weightSums.merge(itemId, sim, Double::sum);
        }
    }
    // 归一化
    for (Long itemId : predictions.keySet()) {
        double weightSum = weightSums.getOrDefault(itemId, 0.0);
        if (weightSum > 0) {
            result.put(itemId, predictions.get(itemId) / weightSum);
        }
    }
    return result;
}""")
    
    doc.add_heading('4.3.4 热门物品回退', level=3)
    
    add_formatted_paragraph(doc,
        '当无法找到有效邻居或预测结果为空时，算法回退到热门物品推荐。热门度评分综合考虑物品平均分和评分数量，公式为：popularityScore = (avgRating - 1.0) / 4.0，确保评分归一化到[0,1]区间。',
        font_size=12, space_after=6)
    
    # ========== 4.4 基于物品的协同过滤算法 ==========
    doc.add_heading('4.4 基于物品的协同过滤算法（Item-Based CF）', level=2)
    
    add_formatted_paragraph(doc,
        'ItemBasedCF类实现了基于物品的协同过滤算法。与User-CF不同，Item-CF通过计算物品间的相似度来推荐与用户已评分物品相似的其他物品。算法支持两种调用方式：直接传入userItem矩阵或同时传入预计算的itemUsers矩阵以提升性能。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.4.1 核心参数', level=3)
    
    headers = ['参数', '值', '说明']
    rows = [
        ['MIN_SIMILARITY', '0.0', '最小相似度阈值'],
        ['TOP_K_SIMILAR_ITEMS', '50', '最大相似物品数量'],
        ['MIN_OVERLAP', '2', '最小共同用户数'],
    ]
    add_table_with_data(doc, headers, rows)
    doc.add_paragraph()
    
    doc.add_heading('4.4.2 物品-用户矩阵构建', level=3)
    
    add_formatted_paragraph(doc,
        'Item-CF需要物品到用户的反向映射。buildItemUsers方法从userItem矩阵构建itemUsers矩阵，将"用户→物品评分"转换为"物品→用户评分"。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private Map<Long, Map<Long, Double>> buildItemUsers(
        Map<Long, Map<Long, Double>> userItem) {
    Map<Long, Map<Long, Double>> itemUsers = new HashMap<>();
    for (Map.Entry<Long, Map<Long, Double>> userEntry : userItem.entrySet()) {
        Long userId = userEntry.getKey();
        for (Map.Entry<Long, Double> ratingEntry : userEntry.getValue().entrySet()) {
            itemUsers.computeIfAbsent(ratingEntry.getKey(), k -> new HashMap<>())
                     .put(userId, ratingEntry.getValue());
        }
    }
    return itemUsers;
}""")
    
    doc.add_heading('4.4.3 评分预测', level=3)
    
    add_formatted_paragraph(doc,
        'Item-CF的评分预测遍历用户已评分的每个物品，计算该物品与所有候选物品的相似度，以相似度为权重对候选物品进行加权评分。对于每对物品，需要计算共同用户数，只有共同用户数达到MIN_OVERLAP才计算相似度。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private Map<Long, Double> predictRatings(Map<Long, Double> targetRatings,
                                         Map<Long, Map<Long, Double>> itemUsers) {
    for (Map.Entry<Long, Double> ratedEntry : targetRatings.entrySet()) {
        Long ratedItemId = ratedEntry.getKey();
        double userRating = ratedEntry.getValue();
        
        Map<Long, Double> ratedItemUsers = itemUsers.get(ratedItemId);
        
        for (Map.Entry<Long, Map<Long, Double>> entry : itemUsers.entrySet()) {
            Long candidateItemId = entry.getKey();
            if (targetRatings.containsKey(candidateItemId)) continue;
            
            // 计算共同用户数
            int overlap = countOverlap(ratedItemUsers, entry.getValue());
            if (overlap < MIN_OVERLAP) continue;
            
            // 计算物品间余弦相似度
            double similarity = cosineSimilarity(ratedItemUsers, entry.getValue());
            if (similarity <= MIN_SIMILARITY) continue;
            
            predictions.merge(candidateItemId, similarity * userRating, Double::sum);
            weightSums.merge(candidateItemId, similarity, Double::sum);
        }
    }
    // 归一化
    return normalize(predictions, weightSums);
}""")
    
    # ========== 4.5 基于行为的推荐算法 ==========
    doc.add_heading('4.5 基于行为的推荐算法（Behavior-Based）', level=2)
    
    add_formatted_paragraph(doc,
        'Behavior-Based算法将显式评分转换为隐式行为强度，综合考虑评分值和时效性两个维度。该方法适用于处理用户浏览、点击、收藏等隐式反馈场景。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.5.1 隐式行为矩阵构建', level=3)
    
    add_formatted_paragraph(doc,
        '隐式行为强度计算公式为：strength = (0.2 + 0.8 × score/5.0) × (0.4 + 0.6 × decay)，其中score为原始评分（归一化到[0,5]），decay为时间衰减权重。对于近期交互（decay≥0.8），额外施加15%的放大系数。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private Map<Long, Map<Long, Double>> buildImplicitBehaviorMatrix(
        Map<Long, Map<Long, Double>> matrix,
        Map<Long, Map<Long, Double>> decayMap) {
    for (Map.Entry<Long, Map<Long, Double>> userEntry : matrix.entrySet()) {
        Long userId = userEntry.getKey();
        Map<Long, Double> decays = decayMap.getOrDefault(userId, emptyMap());
        
        for (Map.Entry<Long, Double> itemEntry : userEntry.getValue().entrySet()) {
            double score = clamp(itemEntry.getValue(), 0.0, 5.0);
            double decay = decays.getOrDefault(itemEntry.getKey(), 1.0);
            
            // 行为强度 = 评分基础 × 时效性因子
            double base = 0.2 + 0.8 * (score / 5.0);
            double recencyFactor = 0.4 + 0.6 * decay;
            if (decay >= 0.8) base *= 1.15;  // 近期交互小幅放大
            
            double strength = base * recencyFactor;
            row.put(itemEntry.getKey(), strength);
        }
    }
    return implicit;
}""")
    
    add_formatted_paragraph(doc,
        '构建隐式行为矩阵后，系统调用ItemBasedCF算法进行推荐，将隐式行为强度作为输入评分。',
        font_size=12, space_after=6)
    
    # ========== 4.6 混合推荐算法 ==========
    doc.add_heading('4.6 混合推荐算法（Hybrid）', level=2)
    
    add_formatted_paragraph(doc,
        '混合推荐算法是本系统的核心创新点，融合了五种推荐信号，并根据用户活跃度动态调整各信号权重。混合策略通过sigmoid函数实现权重的平滑过渡，避免了分段函数的边界跳变问题。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.6.1 五种推荐信号', level=3)
    
    headers = ['信号', '来源', '计算方式']
    rows = [
        ['Item-CF', '基于物品的协同过滤', '物品相似度加权评分'],
        ['User-CF', '基于用户的协同过滤', '用户相似度加权评分'],
        ['Popularity', '热门物品', '平均分×log(1+评分次数)'],
        ['Association', '物品关联', '预计算共现相似度'],
        ['Content', '内容相似度', '类别偏好匹配度'],
    ]
    add_table_with_data(doc, headers, rows)
    doc.add_paragraph()
    
    doc.add_heading('4.6.2 动态权重调整', level=3)
    
    add_formatted_paragraph(doc,
        '系统使用sigmoid函数根据用户历史评分数量动态计算各信号权重。归一化活跃度为min(1.0, ratedCount/30.0)，sigmoid中心点设在0.4，斜率为10。当用户活跃度低时，侧重内容相似度和热门物品；当用户活跃度高时，侧重协同过滤信号。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private HybridWeights dynamicWeights(int ratedCount) {
    double normalizedActivity = min(1.0, ratedCount / 30.0);
    double sigmoid = 1.0 / (1.0 + exp(-10.0 * (normalizedActivity - 0.4)));
    
    double itemCfWeight = 0.30 + 0.15 * sigmoid;     // 0.30~0.45
    double userCfWeight = 0.15 + 0.15 * sigmoid;     // 0.15~0.30
    double popularityWeight = 0.25 - 0.15 * sigmoid; // 0.25~0.10
    double associationWeight = 0.10 + 0.02 * sigmoid;// 0.10~0.12
    double contentWeight = 0.20 - 0.15 * sigmoid;    // 0.20~0.05
    
    // 归一化确保权重之和为1
    double total = itemCfWeight + userCfWeight + popularityWeight 
                 + associationWeight + contentWeight;
    return new HybridWeights(itemCfWeight/total, userCfWeight/total, ...);
}""")
    
    doc.add_heading('4.6.3 偏好类别提升', level=3)
    
    add_formatted_paragraph(doc,
        '系统识别用户偏好的类别（平均分≥4.0且加权评分数≥2），计算类别强度用于提升该类别物品的推荐得分。类别强度公式为：strength = min(1.0, ((avg-4.0)×0.7) + min(0.3, (count-2.0)×0.08))。最终推荐得分乘以(1.0 + 0.25 × categoryStrength)进行提升。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""// 混合推荐得分计算
double score = weights.itemCf() * itemRankScore
             + weights.userCf() * userRankScore
             + weights.popularity() * popScore
             + weights.association() * associationScore
             + weights.content() * contentSimilarityScore;
// 偏好类别提升
score *= (1.0 + 0.25 * preferredCategoryBoost);""")
    
    # ========== 4.7 物品关联预计算 ==========
    doc.add_heading('4.7 物品关联预计算', level=2)
    
    add_formatted_paragraph(doc,
        'ItemAssociationPrecomputeService负责离线预计算物品间的共现相似度，在线阶段只做轻量读取与融合。该服务在应用启动时构建快照，之后每15分钟定时刷新，数据变更时标记为脏数据等待下次刷新。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.7.1 共现相似度计算', level=3)
    
    add_formatted_paragraph(doc,
        '物品间相似度基于共同用户数计算，使用余弦相似度公式：sim(i,j) = co_count(i,j) / sqrt(userCount(i) × userCount(j))。每个物品保留相似度最高的前N个邻居（默认120个），使用PriorityQueue高效维护Top-N。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""// 共现计数
for (Set<Long> items : userItems.values()) {
    for (int i = 0; i < list.size(); i++) {
        for (int j = i + 1; j < list.size(); j++) {
            coCount.merge(list.get(i), list.get(j), 1, Integer::sum);
        }
    }
}

// 相似度计算与Top-N维护
double sim = coCount / sqrt(itemUsers * neighborUsers);
PriorityQueue<Map.Entry<Long, Double>> top = new PriorityQueue<>(...);
if (top.size() < neighborLimit) top.offer(entry);
else if (top.peek().getValue() < sim) { top.poll(); top.offer(entry); }""")
    
    # ========== 4.8 时间衰减机制 ==========
    doc.add_heading('4.8 时间衰减机制', level=2)
    
    add_formatted_paragraph(doc,
        '系统引入时间衰减机制，使近期行为对推荐结果的影响更大。衰减权重使用半衰期30天的指数衰减函数：decay = 0.5^(days/30)。该机制主要应用于User-CF算法和偏好类别计算中。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private double decayWeight(Instant now, Instant ratedAt) {
    if (ratedAt == null) return 1.0;
    long days = max(0L, Duration.between(ratedAt, now).toDays());
    return pow(0.5, (double) days / 30.0);  // 半衰期30天
}""")
    
    add_formatted_paragraph(doc, '衰减权重在各模块中的应用：', font_size=12, space_after=6, bold=True)
    
    headers = ['应用场景', '衰减方式', '效果']
    rows = [
        ['User-CF评分矩阵', '评分×decay', '降低历史评分的影响'],
        ['偏好类别计算', '评分×decay后求和', '偏好随时间演变'],
        ['隐式行为强度', '(0.4+0.6×decay)', '近期行为权重更高'],
        ['热门物品排序', '不应用衰减', '保持全站热度稳定性'],
    ]
    add_table_with_data(doc, headers, rows)
    doc.add_paragraph()
    
    # ========== 4.9 冷启动解决方案 ==========
    doc.add_heading('4.9 冷启动解决方案', level=2)
    
    add_formatted_paragraph(doc,
        '系统设计了多层级的冷启动回退策略，确保在推荐结果不足时仍能提供有意义的推荐。回退策略按优先级依次为：主要算法推荐→热门物品回退→目录级冷启动补齐。',
        font_size=12, space_after=6)
    
    doc.add_heading('4.9.1 热门物品回退（改进版）', level=3)
    
    add_formatted_paragraph(doc,
        '热门物品回退在基础热门度排序之上引入了三项改进：随机扰动因子（±10%）避免热门物品固化；类别多样性约束（同类别不超过40%）提升推荐多样性；时间衰减权重使热门度有一定波动空间。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private List<Recommendation> popularFallback(RecommendationContext ctx, 
                                              int need, Set<Long> excluded) {
    Random random = new Random();
    int maxSameCategory = ceil(need * 0.4);  // 同类别上限40%
    
    for (Recommendation rec : ctx.popularityRanking()) {
        if (excluded.contains(rec.getItemId())) continue;
        
        String category = categoryMap.getOrDefault(rec.getItemId(), "");
        if (categoryCount.getOrDefault(category, 0) >= maxSameCategory) continue;
        
        // 随机扰动 ±10%
        double randomFactor = 0.9 + random.nextDouble() * 0.2;
        double adjustedScore = rec.getScore() * randomFactor;
        
        result.add(new Recommendation(rec.getItemId(), adjustedScore));
        if (result.size() >= need) break;
    }
    return result;
}""")
    
    doc.add_heading('4.9.2 目录级冷启动补齐', level=3)
    
    add_formatted_paragraph(doc,
        '当热门物品仍不足时，系统从全量商品目录中补齐。补齐策略动态调整偏好类别权重，根据用户活跃度自适应：activityFactor = min(1.0, userActivity/10.0)，基础权重baseWeight = 0.35 + 0.15 × activityFactor。活跃用户（评分数>5）额外获得0.10的偏好提升。同类别物品数量限制为need的50%。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private List<Recommendation> catalogFallback(..., int need, Set<Long> excluded) {
    int userActivity = userRatings.size();
    double activityFactor = min(1.0, userActivity / 10.0);
    double baseWeight = 0.35 + 0.15 * activityFactor;
    double activityBonus = userActivity > 5 ? 0.10 : 0.0;
    
    for (Item item : allItems) {
        if (excluded.contains(item.getId())) continue;
        
        // 类别多样性约束
        if (categoryCount.getOrDefault(category, 0) >= ceil(need * 0.5)) continue;
        
        // 偏好分数归一化
        double prefScore = categoryPref.get(category) / maxPref;
        double score = 0.05 + baseWeight * prefScore + activityBonus 
                     + 1.0 / (1e9 + item.getId());
        candidates.add(new Recommendation(item.getId(), score));
    }
    return candidates.stream().sorted().limit(need).toList();
}""")
    
    # ========== 4.10 多样化推荐（MMR算法） ==========
    doc.add_heading('4.10 多样化推荐（MMR算法）', level=2)
    
    add_formatted_paragraph(doc,
        '系统实现了MMR（Maximal Marginal Relevance）算法来平衡推荐结果的相关性与多样性。MMR公式为：MMR = λ × relevance - (1-λ) × max(similarity_to_selected)，其中λ=0.7控制相关性与多样性的平衡。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private List<Recommendation> diversifyRecommendations(
        List<Recommendation> ranked, int topN, Map<Long, String> categoryMap) {
    double lambda = 0.7;
    List<Recommendation> selected = new ArrayList<>();
    
    // 选择得分最高的第一个
    selected.add(pool.stream().max(comparingDouble(Rec::getScore)).get());
    
    while (!pool.isEmpty() && selected.size() < topN) {
        Recommendation best = null;
        double bestMmr = NEGATIVE_INFINITY;
        
        for (Recommendation candidate : pool) {
            double relevance = candidate.getScore();
            
            // 计算与已选物品的最大相似度（基于类别）
            double maxSimilarity = 0.0;
            for (Recommendation sel : selected) {
                double sim = sameCategory(candidate, sel) ? 0.8 : 0.1;
                maxSimilarity = max(maxSimilarity, sim);
            }
            
            // 位置权重：越靠前的推荐，多样性要求越低
            double positionWeight = 1.0 - 0.1 * min(selected.size(), 5);
            double mmr = lambda * relevance * positionWeight 
                       - (1.0 - lambda) * maxSimilarity;
            
            if (mmr > bestMmr) { bestMmr = mmr; best = candidate; }
        }
        selected.add(best);
        pool.remove(best);
    }
    return selected;
}""")
    
    add_formatted_paragraph(doc,
        '系统还提供了可调节lambda参数的多样化方法recommendWithDiversity，允许外部调用者根据需求调整多样性级别（0.0-1.0）。lambda参数计算公式为：lambda = 1.0 - diversityLevel × 0.4，范围为0.6-1.0。',
        font_size=12, space_after=6)
    
    # ========== 4.11 推荐理由生成 ==========
    doc.add_heading('4.11 推荐理由生成', level=2)
    
    add_formatted_paragraph(doc,
        '系统为每条推荐生成人性化的解释说明，推荐理由根据算法类型、类别匹配情况和推荐得分动态生成。',
        font_size=12, space_after=6)
    
    headers = ['算法类型', '条件', '推荐理由']
    rows = [
        ['HYBRID', '类别匹配且偏好强度≥1.2', '混合推荐：你最近对X类目偏好明显...'],
        ['HYBRID', '得分≥0.7', '混合推荐：结合了相似用户、相似商品和全站热度...'],
        ['HYBRID', '其他', '混合推荐：综合多种信号后的候选'],
        ['BEHAVIOR_BASED', '类别匹配', '行为推荐：基于近期行为轨迹...'],
        ['ITEM_BASED', '默认', '评分推荐：和你高分评价过的商品相似度较高'],
        ['USER_BASED', '无历史评分', '评分推荐：基于相似用户偏好的冷启动候选'],
        ['USER_BASED', '有历史评分', '评分推荐：与你评分模式接近的用户也偏好...'],
    ]
    add_table_with_data(doc, headers, rows)
    doc.add_paragraph()
    
    # ========== 4.12 推荐上下文缓存 ==========
    doc.add_heading('4.12 推荐上下文缓存', level=2)
    
    add_formatted_paragraph(doc,
        'RecommendationContext内部类用于缓存推荐过程中的中间计算结果，避免重复计算。采用懒加载策略，只在首次访问时计算。缓存内容包括：用户-物品评分矩阵、用户-物品衰减权重矩阵、物品-用户评分矩阵、热门物品排名、物品类别映射。',
        font_size=12, space_after=6)
    
    add_code_block(doc,
"""private final class RecommendationContext {
    private Map<Long, Map<Long, Double>> userItemMatrix;
    private Map<Long, Map<Long, Double>> userItemDecayMap;
    private Map<Long, Map<Long, Double>> itemUserRatingMatrix;
    private List<Recommendation> popularityRanking;
    private final Map<Long, String> categoryCache = new HashMap<>();
    
    Map<Long, Map<Long, Double>> userItemMatrix() {
        if (userItemMatrix == null) {
            userItemMatrix = buildUserItemMatrix();  // 懒加载
        }
        return userItemMatrix;
    }
    
    List<Recommendation> popularityRanking() {
        if (popularityRanking == null) {
            // 热门度 = 平均分 × log(1 + 评分次数)
            for (ItemPopularityStatView row : ratingRepository.findItemPopularityStats()) {
                popularityRanking.add(new Recommendation(
                    row.getItemId(), row.getAvgScore() * log1p(row.getRatingCount())));
            }
            popularityRanking = popularityRanking.stream().sorted().toList();
        }
        return popularityRanking;
    }
}""")
    
    # ========== 4.13 算法流程总结 ==========
    doc.add_heading('4.13 算法流程总结', level=2)
    
    add_formatted_paragraph(doc,
        '完整的推荐流程如下：', font_size=12, space_after=6, bold=True)
    
    add_formatted_paragraph(doc,
        '(1) 接收推荐请求，获取用户ID、推荐数量和算法类型参数。\n'
        '(2) 构建推荐上下文，懒加载用户-物品评分矩阵和衰减权重矩阵。\n'
        '(3) 根据算法类型执行相应的推荐策略：\n'
        '    - USER_BASED：对评分矩阵应用时效衰减后调用UserBasedCF\n'
        '    - ITEM_BASED：调用ItemBasedCF，传入预计算的物品-用户矩阵\n'
        '    - BEHAVIOR_BASED：构建隐式行为矩阵后调用ItemBasedCF\n'
        '    - HYBRID：融合五种推荐信号，应用动态权重和偏好类别提升\n'
        '(4) 合并推荐结果并应用热门物品回退，确保推荐数量达到要求。\n'
        '(5) 对于HYBRID算法，额外应用MMR多样化优化。\n'
        '(6) 生成推荐理由，返回最终推荐结果列表。',
        font_size=11, space_after=6)
    
    # 保存文件
    output_path = r'D:\app\ks\recommendtwo\docs\基于协同过滤的商品推荐系统设计与实现_更新第四章.docx'
    doc.save(output_path)
    print(f'文件已保存至: {output_path}')

if __name__ == '__main__':
    main()
